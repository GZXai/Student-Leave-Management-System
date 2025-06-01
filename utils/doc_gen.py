from docx import Document
from datetime import datetime
import os

UPLOAD_FOLDER = 'generated_docs'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


def generate_leave_doc(request):
    doc = Document()
    doc.add_heading('学生请假条', 0)

    content = [
        f"学生ID: {request.student_id}",
        f"请假类型: {'病假' if request.leave_type == 'sick' else '事假' if request.leave_type == 'personal' else '公假'}",
        f"请假时间: {request.start_date.strftime('%Y-%m-%d')}至{request.end_date.strftime('%Y-%m-%d')}",
        f"请假理由: {request.reason}",
        f"审批状态: {request.status.upper()}",
        f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ]

    for item in content:
        doc.add_paragraph(item)

    filename = f'leave_{request.id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    doc.save(filepath)
    return filepath